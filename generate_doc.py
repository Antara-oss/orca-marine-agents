import subprocess
import sys

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

COLOR_PRIMARY = RGBColor(15, 76, 129)
COLOR_TEXT = RGBColor(40, 40, 40)

def add_title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

def add_subtitle(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(70, 70, 70)

def add_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLOR_PRIMARY

def add_body(text, prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if prefix:
        r_pre = p.add_run(prefix)
        r_pre.font.name = "Calibri"
        r_pre.font.size = Pt(11)
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_TEXT
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT

add_title("PROJECT ORCA")
add_subtitle("Autonomous Satellite Telemetry & Marine Disaster Sentry System")

add_heading("1. Executive Summary")
add_body(
    "Project ORCA is an edge-native, multi-agent artificial intelligence telemetry framework "
    "designed to monitor, cross-verify, and broadcast operational alerts for rapid marine ecological "
    "disasters across Indian coastal waters. Primarily targeting Harmful Algal Blooms (HABs), severe "
    "coastal upwelling events, and benthic hypoxia/dead-zone collapses, ORCA replaces manual, retrospective "
    "satellite imagery inspection with an autonomous, consensus-driven deliberation pipeline."
)

add_heading("2. Operational Challenges & Core Innovation")
add_body(
    "Earth observation satellite swaths (NetCDF/HDF matrices) span hundreds of megabytes per tile. "
    "Transmitting heavy multi-spectral rasters across tactical maritime communication links directly to "
    "vessels at sea is bandwidth-prohibitive.",
    prefix="The Bandwidth Constraint: "
)
add_body(
    "Conventional satellite detection pipelines rely on centralized processing and manual inspection, "
    "often producing advisories 24 to 72 hours after bloom onset. This latency gap results in severe "
    "economic losses for coastal fisheries and high mortality in mariculture zones.",
    prefix="The Latency Gap: "
)
add_body(
    "ORCA executes statistical anomaly isolation directly on the observation swath at the edge, "
    "distilling high-dimensional spatial data into targeted telemetry frames. Autonomous domain-specialized "
    "agents then evaluate the hydrodynamic and biochemical viability of the event, distilling the consensus "
    "down to an actionable 240-bit NavIC hex payload for instantaneous broadcast.",
    prefix="The ORCA Solution: "
)

add_heading("3. Distributed 4-Node Agentic Pipeline")

table = doc.add_table(rows=5, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Agent Node", "Domain Specialization", "Operational Responsibilities"]
col_widths = [Inches(1.8), Inches(1.8), Inches(3.0)]

for i, title in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = title
    p = cell.paragraphs[0]
    r = p.runs[0]
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.bold = True

agent_data = [
    (
        "Node 1: Orbital Perception Sentry",
        "Edge Spatial Observation",
        "Ingests multi-spectral rasters (SST, Chl-a, winds); executes statistical Z-score filtering; isolates significant anomaly coordinates."
    ),
    (
        "Node 2: Hydrodynamic Oceanographer",
        "Physical Fluid Dynamics",
        "Assesses wind stress curl, thermocline shoaling, and offshore Ekman transport divergence to verify upwelling dynamics."
    ),
    (
        "Node 3: Marine Biogeochemist",
        "Biological Oceanography",
        "Diagnoses suspected taxa (e.g., Skeletonema costatum, Noctiluca); forecasts Biochemical Oxygen Demand (BOD) and hypoxia progression."
    ),
    (
        "Node 4: Tactical Synthesizer",
        "Command & Disaster Comms",
        "Harmonizes node analyses into an operational alert tier; generates targeted directives; formats a compressed 240-bit NavIC S-band hex payload."
    )
]

for row_idx, data_row in enumerate(agent_data, start=1):
    for col_idx, text_val in enumerate(data_row):
        cell = table.cell(row_idx, col_idx)
        cell.text = text_val
        p = cell.paragraphs[0]
        p.paragraph_format.line_spacing = 1.15
        r = p.runs[0]
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        if col_idx == 0:
            r.font.bold = True

for row in table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

add_heading("4. Software Architecture & Technical Stack")
add_body("Python 3.11+ environment with strict typed schemas enforced via Pydantic.", prefix="Core Language & Schemas: ")
add_body("Google GenAI SDK leveraging Gemini 3.6 Flash configured with structured JSON schema outputs.", prefix="Agent Reasoning Core: ")
add_body("xarray, netCDF4, and numpy for multi-dimensional spatial grid manipulation.", prefix="Geospatial Telemetry: ")
add_body("Streamlit frontend integrated with Folium for interactive coastal risk mapping.", prefix="Interactive Dashboard: ")

add_heading("5. Implementation Roadmap")
add_body("Multi-agent test validated with verified physical and biogeochemical outputs.", prefix="Milestone 1 (Complete): ")
add_body("NetCDF telemetry array extraction pipeline feeding structured anomaly vectors to agents.", prefix="Milestone 2 (Current): ")
add_body("Live operational UI visualizing real-time agent deliberation traces and NavIC payloads.", prefix="Milestone 3 (Next): ")
add_body("Alignment with high-performance computing hardware for large-scale spatial raster training.", prefix="Milestone 4 (Mentorship Scope): ")

out_path = "Project_ORCA_Architecture_Brief.docx"
doc.save(out_path)
print(f"Successfully generated {out_path}")
